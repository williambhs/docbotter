from docx import Document

def clear_document(document_path):
    try:
        doc = Document(document_path)

        # Clearing pockets
        for paragraph in doc.paragraphs:
            paragraph.text = ""
            if paragraph.style.name == 'Heading 1':  # Verbatim uses 'Heading1' as the style for pockets
                paragraph.style = doc.styles['Normal']

        # Clearing tables 
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = ""
                    cell.paragraphs[0].style = doc.styles['Normal']

        # Clearing shapes
        for shape in doc.inline_shapes:
            if shape.has_text_frame:
                shape.text_frame.text = ""
                shape.text_frame.paragraphs[0].style = doc.styles['Normal']

        # Clearing headers
        for section in doc.sections:
            for header in section.header.paragraphs:
                header.text = ""
                if header.style.name == 'Heading 1':  # Adjust the style name as needed
                    header.style = doc.styles['Normal']

        # Clearing footers
        for section in doc.sections:
            for footer in section.footer.paragraphs:
                footer.text = ""
                if footer.style.name == 'Heading 1':  # Adjust the style name as needed
                    footer.style = doc.styles['Normal']

        doc.save(document_path)
    except FileNotFoundError:
        raise FileNotFoundError("Specified file does not exist.")
def list_document_elements(document):
    try:
        doc = Document(document)

        # List paragraphs
        print("Paragraphs:")
        for i, paragraph in enumerate(doc.paragraphs, 1):
            print(f"{i}. {paragraph.text}")

        # List tables
        print("\nTables:")
        for i, table in enumerate(doc.tables, 1):
            print(f"{i}. Table with {len(table.rows)} rows and {len(table.columns)} columns")

        # List text in shapes (e.g., images)
        print("\nShapes/Imgs:")
        for i, shape in enumerate(doc.inline_shapes, 1):
            if shape.has_text_frame:
                print(f"{i}. Shape Text: {shape.text_frame.text}")

        # List headers
        print("\nHeaders:")
        for i, section in enumerate(doc.sections, 1):
            for j, header in enumerate(section.header.paragraphs, 1):
                print(f"{i}.{j}. {header.text}")

        # List footers
        print("\nFooters:")
        for i, section in enumerate(doc.sections, 1):
            for j, footer in enumerate(section.footer.paragraphs, 1):
                print(f"{i}.{j}. {footer.text}")

    except FileNotFoundError:
        raise FileNotFoundError("Specified file does not exist.")
def print_xml(document):
    try:
        doc = Document(document)

        # Access the underlying XML
        xml_content = doc.element.xml

        # Print the XML content
        print(xml_content)

    except FileNotFoundError:
        raise FileNotFoundError("Specified file does not exist.")
    
# Testing--------------------

document_path = 'docxtest.docx'
#print_xml(document_path)
clear_document(document_path)
#list_document_elements(document_path)






